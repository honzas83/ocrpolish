import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.section import Section
from docx.shared import Pt

from ocrpolish.data_model import PageMetadata
from ocrpolish.utils.metadata import extract_page_number

# XML 1.0 restricted characters: \x00-\x08, \x0b-\x0c, \x0e-\x1f
ILLEGAL_XML_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def sanitize_xml(text: str) -> str:
    """Remove characters that are not allowed in XML 1.0."""
    if not text:
        return ""
    return ILLEGAL_XML_CHARS_RE.sub("", text)


def split_markdown_by_pages(content: str, scan_paragraphs: int = 3) -> list[PageMetadata]:
    """
    Split markdown content into pages based on '---' and '# Page X' markers.
    Extracts page metadata (page numbers, potential headers/footers).
    """
    # Normalize line endings and split by lines
    lines = content.splitlines()
    raw_pages: list[list[str]] = []
    current_page_lines: list[str] = []

    for line in lines:
        trimmed = line.strip()
        # Detect page boundaries: '---' or '# Page X'
        if trimmed == "---" or re.match(r"^#\s*Page\s+\d+$", trimmed, re.IGNORECASE):
            if current_page_lines:
                raw_pages.append(current_page_lines)
                current_page_lines = []
            continue

        current_page_lines.append(line)

    if current_page_lines:
        raw_pages.append(current_page_lines)

    processed_pages: list[PageMetadata] = []
    for page_lines in raw_pages:
        if not any(line.strip() for line in page_lines):
            continue

        metadata = PageMetadata()

        # 1. Extract page number and remove from body if found
        # We remove ALL instances found on the page to prevent duplication in body
        remaining_lines = []
        for line in page_lines:
            p_num = extract_page_number(line)
            if p_num is not None:
                if metadata.page_number is None:
                    metadata.page_number = p_num
            else:
                remaining_lines.append(line)

        if not remaining_lines:
            continue

        # 2. Identify candidates for repeated headers/footers
        content_lines = [line for line in remaining_lines if line.strip()]

        if content_lines:
            metadata.header_candidates = content_lines[:scan_paragraphs]
            metadata.footer_candidates = content_lines[-scan_paragraphs:]
            metadata.body_lines = remaining_lines

        processed_pages.append(metadata)

    return processed_pages


def calculate_font_size(text: str, base_size: int = 10, max_lines: int = 50) -> float:
    """
    Heuristic to scale font size so that content fits on a single DOCX page.
    Assumes standard A4/Letter margins.
    """
    lines = text.splitlines()
    line_count = len(lines)

    if line_count <= max_lines:
        return float(base_size)

    # Scale down proportionally if we exceed the target line count
    scale_factor = max_lines / line_count
    # Don't scale below 6pt for readability
    return max(6.0, base_size * scale_factor)


def _add_margin_line(
    container: "Section", # type: ignore
    left: str = "",
    center: str = "",
    right: str = "",
    font_name: str = "Consolas",
    font_size: float = 10.0,
) -> None:
    """Add a line with left, centered, and right aligned components using tab stops."""
    if not any([left, center, right]):
        return

    left, center, right = sanitize_xml(left), sanitize_xml(center), sanitize_xml(right)

    p = container.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.tab_stops.clear_all()
    # Remove spacing between lines in headers/footers
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    # A4 content width is ~523pt (595.3 - 72)
    # Centered at 261.5pt, Right at 523pt
    pf.tab_stops.add_tab_stop(Pt(261.5), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    pf.tab_stops.add_tab_stop(Pt(523), alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Use double tabs to skip positions if needed
    run = p.add_run(f"{left}\t{center}\t{right}")
    run.font.name = font_name
    run.font.size = Pt(font_size)


def _apply_header_footer_metadata(
    section: Section,
    metadata: PageMetadata,
    font_name: str,
    font_size: float,
) -> None:
    """Apply metadata to section headers and footers with centered metadata lines."""
    # 1. Clear Header
    header = section.header
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    # Concatenate all header metadata into one centered line
    all_header_metadata = []
    if metadata.header_left:
        all_header_metadata.extend(metadata.header_left)
    if metadata.header_right:
        all_header_metadata.extend(metadata.header_right)
    header_metadata_str = " ".join(all_header_metadata)

    center_marker = f"- {metadata.original_page_number} -" if metadata.original_page_number else ""

    # Header Line 1: Concatenated Metadata (Centered)
    _add_margin_line(header, center=header_metadata_str, font_name=font_name, font_size=font_size)
    # Header Line 2: Original Page Number (Centered)
    _add_margin_line(header, center=center_marker, font_name=font_name, font_size=font_size)

    # 2. Clear Footer
    footer = section.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    # Concatenate all footer metadata into one centered line
    all_footer_metadata = []
    if metadata.footer_left:
        all_footer_metadata.extend(metadata.footer_left)
    if metadata.footer_right:
        all_footer_metadata.extend(metadata.footer_right)
    footer_metadata_str = " ".join(all_footer_metadata)

    pdf_label = f"PDF Page {metadata.pdf_page_number}" if metadata.pdf_page_number else ""

    # Footer Line 1: Original Page Number (Centered) + PDF Page N (Right)
    _add_margin_line(footer, center=center_marker, right=pdf_label, font_name=font_name, font_size=font_size)
    # Footer Line 2: Concatenated Metadata (Centered)
    _add_margin_line(footer, center=footer_metadata_str, font_name=font_name, font_size=font_size)


def _setup_section_margins(doc: Document) -> None:
    """Adjust margins to optimize space (0.3" top/bottom, 0.5" sides)."""
    for section in doc.sections:
        # Reduced to 0.3 inch
        section.top_margin = section.bottom_margin = Pt(21.6)
        # Keep 0.5 inch
        section.left_margin = section.right_margin = Pt(36)


def _apply_page_number(
    section: "Section",  # type: ignore[name-defined]
    page_num: int,
    font_name: str,
    font_size: float,
) -> None:
    """Apply centered page number to header and footer."""
    text = sanitize_xml(f"- {page_num} -")
    
    # Apply to header
    header = section.header
    for p in header.paragraphs:
        p.text = ""
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)

    # Apply to footer
    footer = section.footer
    for p in footer.paragraphs:
        p.text = ""
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)


def _is_table_start(line: str, next_line: str | None) -> bool:
    """Check if lines indicate the start of a table."""
    trimmed = line.strip()
    if trimmed.startswith("|"):
        return True
    if next_line:
        next_trimmed = next_line.strip()
        if "|" in line and "-" in next_trimmed and ("|" in next_trimmed or "---" in next_trimmed):
            return True
    return False


def _render_text_block(
    doc: Document,
    text_block: list[str],
    font_name: str,
    font_size: float,
) -> None:
    """Render a block of text lines as a single paragraph."""
    if not text_block:
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    run = p.add_run()
    # Manual page breaks removed because we use WD_SECTION.NEW_PAGE

    for j, text_line in enumerate(text_block):
        run.add_text(sanitize_xml(text_line))
        if j < len(text_block) - 1:
            run.add_break()

    run.font.name = font_name
    run.font.size = Pt(font_size)


def create_docx_from_pages(
    pages: list[PageMetadata],
    output_path: Path,
    font_name: str = "Consolas",
    font_size: int = 10,
) -> None:
    """
    Create a DOCX file with optimized performance and page numbers in headers/footers.
    """
    doc = Document()
    # Remove default paragraph to start clean
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Initial setup for first section
    _setup_section_margins(doc)

    # Render content
    for i, page_metadata in enumerate(pages):
        page_content = "\n".join(page_metadata.body_lines).strip()
        dynamic_size = calculate_font_size(page_content, base_size=font_size, max_lines=65)

        # Create/Get section for the page
        if i == 0:
            section = doc.sections[0]
        else:
            # add_section(WD_SECTION.NEW_PAGE) handles the page break automatically
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            section.top_margin = section.bottom_margin = Pt(36)
            section.left_margin = section.right_margin = Pt(36)

        # Apply enhanced header and footer metadata
        _apply_header_footer_metadata(
            section,
            page_metadata,
            font_name,
            dynamic_size,
        )

        lines = page_content.splitlines()
        line_idx = 0
        
        if not lines:
            doc.add_paragraph()
        
        while line_idx < len(lines):
            line = lines[line_idx]
            next_line = lines[line_idx + 1] if line_idx + 1 < len(lines) else None

            if _is_table_start(line, next_line):
                table_lines = []
                while line_idx < len(lines) and "|" in lines[line_idx]:
                    table_lines.append(lines[line_idx])
                    line_idx += 1

                if table_lines:
                    _render_table(doc, table_lines, font_name, dynamic_size)
            else:
                text_block = []
                while line_idx < len(lines):
                    curr = lines[line_idx]
                    nxt = lines[line_idx + 1] if line_idx + 1 < len(lines) else None
                    if _is_table_start(curr, nxt):
                        break
                    text_block.append(curr)
                    line_idx += 1

                if text_block:
                    _render_text_block(
                        doc,
                        text_block,
                        font_name,
                        dynamic_size,
                    )

    doc.save(str(output_path))


def _render_table(doc: Document, table_lines: list[str], font_name: str, font_size: float) -> None:
    """Helper to convert markdown table lines into a native Word table."""
    rows_data = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped:
            continue

        if "-" in stripped:
            remaining = stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
            if len(remaining) == 0:
                continue

        content = stripped[1:] if stripped.startswith("|") else stripped
        content = content[:-1] if content.endswith("|") else content
        cells = [c.strip() for c in content.split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"
    table.autofit = True

    for r_idx, row_data in enumerate(rows_data):
        row_cells = table.rows[r_idx].cells
        for c_idx, cell_value in enumerate(row_data):
            if c_idx < num_cols:
                cell = row_cells[c_idx]
                cell.text = sanitize_xml(cell_value)
                for paragraph in cell.paragraphs:
                    pf = paragraph.paragraph_format
                    pf.line_spacing = 1.0
                    pf.space_before = pf.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
