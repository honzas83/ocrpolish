import sys
from pathlib import Path
from unittest.mock import patch

from docx import Document

from ocrpolish.cli import main


def test_docx_page_number_extraction(tmp_path: Path) -> None:
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    docx_dir = tmp_path / "docx"
    input_dir.mkdir()
    
    # Page numbers: - 1 - and - 2 -
    content = "Page 1 content\n- 1 -\n---\n# Page 2\n- 2 -\nPage 2 content"
    test_file = input_dir / "test.md"
    test_file.write_text(content, encoding="utf-8")
    
    with patch.object(sys, "argv", ["ocrpolish", str(input_dir), str(output_dir), "--docx", str(docx_dir)]):
        main()
        
    docx_file = docx_dir / "test.docx"
    assert docx_file.exists()
    
    doc = Document(str(docx_file))
    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "- 1 -" not in body_text
    assert "- 2 -" not in body_text
    
    # Check headers and footers for all sections
    for section in doc.sections:
        header_text = "\n".join(p.text for p in section.header.paragraphs)
        footer_text = "\n".join(p.text for p in section.footer.paragraphs)
        assert "-" in header_text
        assert "-" in footer_text

def test_docx_recurring_text_remains_in_body(tmp_path: Path) -> None:
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    docx_dir = tmp_path / "docx"
    input_dir.mkdir()
    
    # "SECRET" is repeated at top of 4/5 pages (80%)
    pages = []
    for i in range(5):
        pages.append(f"SECRET\nPage {i+1} content")
    
    content = "\n---\n".join(pages)
    test_file = input_dir / "test.md"
    test_file.write_text(content, encoding="utf-8")
    
    with patch.object(sys, "argv", ["ocrpolish", str(input_dir), str(output_dir), "--docx", str(docx_dir)]):
        main()
        
    docx_file = docx_dir / "test.docx"
    doc = Document(str(docx_file))
    
    body_text = "\n".join(p.text for p in doc.paragraphs)
    # recurring text should NOW remain in body
    assert "SECRET" in body_text
