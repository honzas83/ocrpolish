import sys
from pathlib import Path
from unittest.mock import patch

from docx import Document

from ocrpolish.cli import main


def test_docx_generation_flag(tmp_path: Path) -> None:
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    docx_dir = tmp_path / "docx"
    input_dir.mkdir()

    test_file = input_dir / "test.md"
    test_file.write_text("Page 1\n---\n# Page 2\nPage 2 content", encoding="utf-8")

    # Run CLI with --docx DIR
    with patch.object(
        sys, "argv", ["ocrpolish", str(input_dir), str(output_dir), "--docx", str(docx_dir)]
    ):
        main()

    docx_file = docx_dir / "test.docx"
    assert docx_file.exists()


def test_no_docx_generation_by_default(tmp_path: Path) -> None:
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()

    test_file = input_dir / "test.md"
    test_file.write_text("Some content", encoding="utf-8")

    # Run CLI without --docx
    with patch.object(sys, "argv", ["ocrpolish", str(input_dir), str(output_dir)]):
        main()

    docx_file = output_dir / "test.docx"
    assert not docx_file.exists()


def test_docx_page_content(tmp_path: Path) -> None:
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    docx_dir = tmp_path / "docx"
    input_dir.mkdir()

    test_file = input_dir / "test.md"
    test_file.write_text("P1\n---\n# Page 2\nP2", encoding="utf-8")

    with patch.object(
        sys, "argv", ["ocrpolish", str(input_dir), str(output_dir), "--docx", str(docx_dir)]
    ):
        main()

    docx_file = docx_dir / "test.docx"
    doc = Document(str(docx_file))
    # We expect 2 pages.
    # The first paragraph is P1.
    # The second paragraph (after page break) is P2.
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "P1" in texts
    assert "P2" in texts


def test_docx_sections_and_footers(tmp_path: Path) -> None:
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    docx_dir = tmp_path / "docx"
    input_dir.mkdir()

    test_file = input_dir / "test.md"
    content = "# Page 1\nContent 1\n# Page 2\nContent 2"
    test_file.write_text(content, encoding="utf-8")

    with patch.object(
        sys, "argv", ["ocrpolish", str(input_dir), str(output_dir), "--docx", str(docx_dir)]
    ):
        main()

    docx_file = docx_dir / "test.docx"
    doc = Document(str(docx_file))
    
    # Check sections
    assert len(doc.sections) >= 2
    
    # Check footers (US1)
    for i, section in enumerate(doc.sections):
        footer_text = "".join(p.text for p in section.footer.paragraphs)
        assert f"PDF Page {i+1}" in footer_text
