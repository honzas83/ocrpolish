from pathlib import Path

from docx import Document
from docx.shared import Pt

from ocrpolish.data_model import PageMetadata
from ocrpolish.utils.docx_utils import (
    calculate_font_size,
    create_docx_from_pages,
    sanitize_xml,
)


def test_sanitize_xml() -> None:
    # No illegal characters
    assert sanitize_xml("Hello World") == "Hello World"
    # Illegal control characters (0x00, 0x07, 0x1F)
    # 0x09 (tab), 0x0A (line feed), 0x0D (carriage return) are allowed
    input_text = "Hello\x00World\x07!\x1F\t\n\r"
    expected = "HelloWorld!\t\n\r"
    assert sanitize_xml(input_text) == expected
    # Empty string
    assert sanitize_xml("") == ""
    # None
    assert sanitize_xml(None) == ""


def test_calculate_font_size() -> None:
    # Under limit
    text_short = "Line\n" * 10
    expected_base = 10.0
    assert calculate_font_size(text_short, base_size=10, max_lines=50) == expected_base

    # Over limit
    text_long = "Line\n" * 100
    # 50 / 100 * 10 = 5.0, but capped at 6.0
    expected_min = 6.0
    assert calculate_font_size(text_long, base_size=10, max_lines=50) == expected_min

    # Mid range
    text_mid = "Line\n" * 80
    # 50 / 80 * 10 = 6.25
    expected_mid = 6.25
    assert calculate_font_size(text_mid, base_size=10, max_lines=50) == expected_mid


def test_create_docx_scaling(tmp_path: Path) -> None:
    # One short page (1 line), one very long page (100 lines)
    pages = [
        PageMetadata(pdf_page_number=1, body_lines=["Short page"]),
        PageMetadata(pdf_page_number=2, body_lines=["Long"] * 100)
    ]
    output_path = tmp_path / "scale_test.docx"
    create_docx_from_pages(pages, output_path)

    doc = Document(str(output_path))
    # Each PageMetadata block results in at least one paragraph
    assert len(doc.paragraphs) >= 2

    # First page should be 10pt (first paragraph)
    assert doc.paragraphs[0].runs[0].font.size == Pt(10)
    # The last paragraph should have the scaled font
    last_para = doc.paragraphs[-1]
    # 65 (max_lines) / 100 (actual) * 10 (base) = 6.5
    assert last_para.runs[0].font.size == Pt(6.5)


def test_create_docx_with_illegal_characters(tmp_path: Path) -> None:
    # Pages with illegal XML characters
    pages = [
        PageMetadata(
            pdf_page_number=1, 
            body_lines=["Illegal\x00character", "Allowed\tTab"],
            header_left=["Header\x07Metadata"],
            footer_right=["Footer\x1FMetadata"]
        )
    ]
    output_path = tmp_path / "illegal_char_test.docx"
    
    # Should not raise "All strings must be XML compatible"
    create_docx_from_pages(pages, output_path)
    
    assert output_path.exists()
    doc = Document(str(output_path))
    
    # Verify content (illegal characters stripped)
    # The body lines are joined with \n in create_docx_from_pages
    body_text = doc.paragraphs[0].text
    assert "Illegalcharacter" in body_text
    assert "Allowed\tTab" in body_text
    
    # Verify header/footer
    section = doc.sections[0]
    header_text = section.header.paragraphs[0].text
    assert "HeaderMetadata" in header_text
    
    footer_text = section.footer.paragraphs[1].text
    assert "FooterMetadata" in footer_text


def test_native_table_generation(tmp_path: Path) -> None:
    # We create PageMetadata directly as split_markdown_by_pages is deprecated/changed
    content = ["Header 1 | Header 2", "| --- | --- |", "Row 1 Col 1 | Row 1 Col 2"]
    pages = [PageMetadata(pdf_page_number=1, body_lines=content)]
    output_path = tmp_path / "table_test.docx"
    create_docx_from_pages(pages, output_path)

    assert output_path.exists()
    doc = Document(str(output_path))

    # Verify a table was created
    assert len(doc.tables) == 1
    table = doc.tables[0]
    expected_rows = 2 # Header + Row (separator skipped)
    expected_cols = 2
    assert len(table.rows) == expected_rows
    assert len(table.columns) == expected_cols

    # Check content
    assert table.cell(0, 0).text == "Header 1"
    assert table.cell(1, 1).text == "Row 1 Col 2"

    # Check styling (Consolas)
    assert table.cell(0, 0).paragraphs[0].runs[0].font.name == "Consolas"
